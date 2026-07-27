"""
Module 5: ADMET Profiling, SDF Export & DOCX Report Generation
"""

from __future__ import annotations
import io
import math
import datetime
from typing import List, Dict, Optional, Tuple

from rdkit import Chem
from rdkit.Chem import Descriptors, QED, rdMolDescriptors, AllChem
from rdkit.Chem import SDWriter

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── ADMET descriptor calculation ──────────────────────────────────────────────
def _calc_logs(mol: Chem.Mol) -> float:
    """
    Estimate aqueous solubility (LogS) via the ESOL model (Delaney 2004).
    LogS = 0.16 - 0.63·cLogP - 0.0062·MW + 0.066·RBonds - 0.74·ArFraction
    """
    try:
        mw      = Descriptors.MolWt(mol)
        clogp   = Descriptors.MolLogP(mol)
        rotb    = rdMolDescriptors.CalcNumRotatableBonds(mol)
        n_arom  = sum(1 for a in mol.GetAtoms() if a.GetIsAromatic())
        ar_frac = n_arom / max(mol.GetNumHeavyAtoms(), 1)
        logs = 0.16 - 0.63 * clogp - 0.0062 * mw + 0.066 * rotb - 0.74 * ar_frac
        return round(logs, 3)
    except Exception:
        return float("nan")


def compute_admet_profile(mol: Chem.Mol, name: str = "", dg: float = 0.0) -> Dict:
    """Full ADMET profile for a single hit molecule."""
    try:
        mw      = round(Descriptors.ExactMolWt(mol), 2)
        logp    = round(Descriptors.MolLogP(mol), 2)
        hbd     = rdMolDescriptors.CalcNumHBD(mol)
        hba     = rdMolDescriptors.CalcNumHBA(mol)
        tpsa    = round(rdMolDescriptors.CalcTPSA(mol), 2)
        rotb    = rdMolDescriptors.CalcNumRotatableBonds(mol)
        qed_val = round(QED.qed(mol), 3)
        logs    = _calc_logs(mol)
        rings   = rdMolDescriptors.CalcNumRings(mol)
        arom_r  = rdMolDescriptors.CalcNumAromaticRings(mol)

        # Lipinski violations
        viols = sum([mw > 500, logp > 5, hbd > 5, hba > 10])

        # Drug-likeness label
        if qed_val >= 0.7 and viols == 0:
            drug_like = "Excellent"
        elif qed_val >= 0.5 and viols <= 1:
            drug_like = "Good"
        elif qed_val >= 0.3:
            drug_like = "Moderate"
        else:
            drug_like = "Poor"

        # BBB penetration (simple heuristic)
        bbb = "Yes" if (mw < 400 and logp > 0 and tpsa < 90 and hbd <= 3) else "No"

        return {
            "Name":             name,
            "ΔG (kcal/mol)":   round(dg, 3),
            "MW":               mw,
            "LogP":             logp,
            "HBD":              hbd,
            "HBA":              hba,
            "TPSA":             tpsa,
            "RotBonds":         rotb,
            "QED":              qed_val,
            "LogS":             logs,
            "Rings":            rings,
            "AromaticRings":    arom_r,
            "LipinskiViol":     viols,
            "DrugLikeness":     drug_like,
            "BBBPenetration":   bbb,
        }
    except Exception as e:
        return {"Name": name, "ΔG (kcal/mol)": dg, "Error": str(e)}


def build_admet_dataframe(
    mols:   List[Chem.Mol],
    scores: List[float],
    names:  Optional[List[str]] = None,
) -> pd.DataFrame:
    """Compute ADMET profiles for all hit mols and return as a DataFrame."""
    records = []
    for i, (mol, dg) in enumerate(zip(mols, scores)):
        name = (names[i] if names else None) or mol.GetPropsAsDict().get("_Name", f"hit_{i+1}")
        records.append(compute_admet_profile(mol, name, dg))
    return pd.DataFrame(records)


# ── SDF Export ────────────────────────────────────────────────────────────────
def export_hits_sdf(
    mols:   List[Chem.Mol],
    scores: List[float],
    names:  Optional[List[str]] = None,
    admet_df: Optional[pd.DataFrame] = None,
) -> bytes:
    """
    Write all top-hit 3D conformations into one consolidated SDF file.
    Embeds ΔG, ADMET tags, and rank as mol properties.
    Returns raw bytes.
    """
    buf = io.StringIO()
    writer = SDWriter(buf)

    for rank, (mol, dg) in enumerate(zip(mols, scores), start=1):
        mol = Chem.RWMol(mol)
        name = (names[rank - 1] if names else None) or \
               mol.GetPropsAsDict().get("_Name", f"hit_{rank}")

        # Ensure 3D coords
        if mol.GetNumConformers() == 0:
            mol_h = Chem.AddHs(mol)
            AllChem.EmbedMolecule(mol_h, AllChem.ETKDGv3())
            try:
                AllChem.MMFFOptimizeMolecule(mol_h)
            except Exception:
                pass
            mol = Chem.RemoveHs(mol_h)

        # Set properties
        mol.SetProp("_Name",            name)
        mol.SetDoubleProp("DeltaG_kcal_mol", round(dg, 3))
        mol.SetIntProp("Rank",          rank)
        mol.SetDoubleProp("MW",         round(Descriptors.ExactMolWt(mol), 2))
        mol.SetDoubleProp("LogP",       round(Descriptors.MolLogP(mol), 2))
        mol.SetDoubleProp("TPSA",       round(rdMolDescriptors.CalcTPSA(mol), 2))
        mol.SetDoubleProp("QED",        round(QED.qed(mol), 3))
        mol.SetDoubleProp("LogS_ESOL",  _calc_logs(mol))

        if admet_df is not None and rank - 1 < len(admet_df):
            row = admet_df.iloc[rank - 1]
            mol.SetProp("DrugLikeness",   str(row.get("DrugLikeness", "")))
            mol.SetProp("BBBPenetration", str(row.get("BBBPenetration", "")))
            mol.SetIntProp("LipinskiViol", int(row.get("LipinskiViol", 0)))

        writer.write(mol)

    writer.close()
    return buf.getvalue().encode("utf-8")


# ── DOCX Report ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via direct XML (python-docx limitation)."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def generate_docx_report(
    admet_df:     pd.DataFrame,
    n_input:      int,
    n_after_filter: int,
    top_n:        int,
    protein_info: Optional[Dict] = None,
    grid_info:    Optional[Dict] = None,
) -> bytes:
    """
    Generate a professional DOCX screening report.
    Returns raw bytes of the .docx file.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(section, attr, Inches(1.0))

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("DeepDock-AI Virtual Screening Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)  # deep blue

    doc.add_paragraph(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    summary_data = [
        ("Total ligands uploaded",       n_input),
        ("Passed ADMET pre-filter",       n_after_filter),
        ("Top hits reported",             top_n),
        ("Filter pass rate",              f"{n_after_filter/max(n_input,1)*100:.1f}%"),
    ]

    if not admet_df.empty and "ΔG (kcal/mol)" in admet_df.columns:
        dg_col = admet_df["ΔG (kcal/mol)"]
        summary_data += [
            ("Best ΔG (kcal/mol)",            f"{dg_col.min():.3f}"),
            ("Mean ΔG of top hits (kcal/mol)", f"{dg_col.mean():.3f}"),
            ("Median QED score",              f"{admet_df['QED'].median():.3f}" if "QED" in admet_df else "N/A"),
        ]

    tbl = doc.add_table(rows=len(summary_data) + 1, cols=2)
    tbl.style = "Table Grid"
    headers = ["Metric", "Value"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "0D47A1")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (k, v) in enumerate(summary_data, start=1):
        tbl.rows[i].cells[0].text = str(k)
        tbl.rows[i].cells[1].text = str(v)
        if i % 2 == 0:
            for cell in tbl.rows[i].cells:
                _set_cell_bg(cell, "E3F2FD")

    doc.add_paragraph()

    # ── Binding Site Info ─────────────────────────────────────────────────────
    if protein_info or grid_info:
        doc.add_heading("2. Target & Docking Grid", level=1)
        if protein_info:
            cx, cy, cz = protein_info.get("centroid", (0, 0, 0))
            doc.add_paragraph(
                f"Binding site centroid (from PDB): "
                f"X={cx:.2f}, Y={cy:.2f}, Z={cz:.2f}  "
                f"[source: {protein_info.get('source','N/A')}, "
                f"atoms: {protein_info.get('n_atoms', 0)}]"
            )
        if grid_info:
            cx, cy, cz = grid_info.get("center", (0, 0, 0))
            sx, sy, sz = grid_info.get("size",   (0, 0, 0))
            doc.add_paragraph(
                f"Grid center (conf.txt): X={cx}, Y={cy}, Z={cz}  |  "
                f"Grid size: {sx}×{sy}×{sz} Å"
            )
        doc.add_paragraph()

    # ── Top Hits Table ────────────────────────────────────────────────────────
    doc.add_heading("3. Top Hit Compounds", level=1)

    display_cols = [
        "Name", "ΔG (kcal/mol)", "MW", "LogP", "QED",
        "LogS", "TPSA", "LipinskiViol", "DrugLikeness", "BBBPenetration",
    ]
    cols = [c for c in display_cols if c in admet_df.columns]
    top_df = admet_df.head(min(50, len(admet_df)))

    tbl2 = doc.add_table(rows=len(top_df) + 1, cols=len(cols))
    tbl2.style = "Table Grid"

    # Header row
    for j, col in enumerate(cols):
        cell = tbl2.rows[0].cells[j]
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        _set_cell_bg(cell, "1565C0")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, (_, row) in enumerate(top_df.iterrows(), start=1):
        for j, col in enumerate(cols):
            val = row.get(col, "")
            cell = tbl2.rows[i].cells[j]
            if isinstance(val, float):
                cell.text = f"{val:.3f}" if not math.isnan(val) else "N/A"
            else:
                cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            if i % 2 == 0:
                _set_cell_bg(cell, "E8F5E9")

    doc.add_paragraph()

    # ── ADMET Interpretation ──────────────────────────────────────────────────
    doc.add_heading("4. ADMET Profile Summary", level=1)

    if "DrugLikeness" in admet_df.columns:
        dl_counts = admet_df["DrugLikeness"].value_counts()
        for label, count in dl_counts.items():
            doc.add_paragraph(
                f"  • {label}: {count} compound(s) ({count/len(admet_df)*100:.1f}%)",
                style="List Bullet",
            )
    doc.add_paragraph()

    if "QED" in admet_df.columns:
        doc.add_paragraph(
            f"QED Statistics — Mean: {admet_df['QED'].mean():.3f}  |  "
            f"Std: {admet_df['QED'].std():.3f}  |  "
            f"Max: {admet_df['QED'].max():.3f}"
        )
    if "LogS" in admet_df.columns:
        doc.add_paragraph(
            f"LogS (ESOL) — Mean: {admet_df['LogS'].mean():.3f}  "
            f"(−4 to 0.5 = drug-like range)"
        )

    # ── Methods ───────────────────────────────────────────────────────────────
    doc.add_heading("5. Methods", level=1)
    methods = doc.add_paragraph()
    methods.add_run("Pre-filtering: ").bold = True
    methods.add_run(
        "Lipinski Rule of 5 (MW≤500, LogP≤5, HBD≤5, HBA≤10), "
        "Veber rules (RotBonds≤10, TPSA≤140), and PAINS A/B/C catalog via RDKit. "
        "Parallel execution via Python multiprocessing.\n"
    )
    methods.add_run("Docking model: ").bold = True
    methods.add_run(
        "Graph Neural Network (3× GCNLayer + Attention Pooling) trained to predict "
        "binding free energy ΔG (kcal/mol). Inference with PyTorch FP16 AMP on "
        "CUDA GPU (Kaggle T4/P100). CUDA OOM automatically recovers by halving batch size.\n"
    )
    methods.add_run("ADMET descriptors: ").bold = True
    methods.add_run(
        "QED (Bickerton 2012), TPSA, LogP (Crippen), MW, LogS (ESOL/Delaney), "
        "Lipinski violations — all computed via RDKit.\n"
    )

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph("DeepDock-AI  |  PyTorch + RDKit Virtual Screening Platform")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x78, 0x78, 0x78)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
