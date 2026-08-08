"""
Stage 3 — Export helpers: CSV, PDB-complex ZIP, DOCX report.

Public interface (matches app.py imports):
    export_csv(df)                                  -> bytes
    export_pdb_complexes_zip(complexes, names, scores) -> bytes
    generate_docx_report(...)                        -> bytes
"""

from __future__ import annotations
import io
import zipfile
from datetime import datetime
from typing import List, Optional, Dict

import pandas as pd


# ── CSV ─────────────────────────────────────────────────────────────────
def export_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


# ── PDB complexes ZIP ──────────────────────────────────────────────────
def export_pdb_complexes_zip(
    complexes: List[str],
    names: List[str],
    scores: List[float],
) -> bytes:
    """
    Packs a list of complex PDB text blocks into a single downloadable ZIP.
    One file per ligand: {rank}_{name}_{score}.pdb
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for i, (pdb_text, name, score) in enumerate(zip(complexes, names, scores)):
            if not pdb_text:
                continue
            safe_name = "".join(c if c.isalnum() or c in "_-" else "_" for c in str(name))
            fname = f"{i + 1:03d}_{safe_name}_{score:.2f}.pdb"
            zf.writestr(fname, pdb_text)
    buf.seek(0)
    return buf.getvalue()


# ── DOCX report ─────────────────────────────────────────────────────────
def generate_docx_report(
    filter_df: Optional[pd.DataFrame],
    docking_df: Optional[pd.DataFrame],
    admet_df: Optional[pd.DataFrame],
    admet_source: str,
    n_input: int,
    n_filtered: int,
    n_docked: int,
    protein_info: Optional[Dict],
    grid_info: Optional[Dict],
    top_n: int,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("DeepDock-AI — Virtual Screening Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True

    doc.add_heading("Pipeline Summary", level=1)
    summary_table = doc.add_table(rows=0, cols=2)
    summary_table.style = "Light Grid Accent 1"
    rows = [
        ("Input ligands", str(n_input)),
        ("Passed ADMET pre-filter", str(n_filtered)),
        ("Successfully docked", str(n_docked)),
        ("ADMET source (Stage 3)", admet_source),
        ("Top-N carried to ADMET", str(top_n)),
    ]
    if protein_info:
        cx, cy, cz = protein_info.get("centroid", (0, 0, 0))
        rows.append(("Binding-site centroid", f"X={cx:.2f}, Y={cy:.2f}, Z={cz:.2f} ({protein_info.get('source', 'n/a')})"))
    if grid_info:
        sx, sy, sz = grid_info.get("size", (20, 20, 20))
        rows.append(("Grid box size (Å)", f"{sx} × {sy} × {sz}"))

    for label, value in rows:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    def _add_df_table(heading: str, df: Optional[pd.DataFrame], max_rows: int = 50):
        doc.add_heading(heading, level=1)
        if df is None or df.empty:
            doc.add_paragraph("No data available for this stage.")
            return
        display_df = df.head(max_rows)
        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(display_df.columns):
            hdr_cells[i].text = str(col)
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = "" if pd.isna(val) else str(val)
        if len(df) > max_rows:
            doc.add_paragraph(f"… {len(df) - max_rows} additional rows omitted for brevity.")

    _add_df_table("Stage 1 — ADMET Pre-Filter Results", filter_df)
    _add_df_table("Stage 2 — Docking Results", docking_df)
    _add_df_table("Stage 3 — ADMET Profiles", admet_df)

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Ligands were pre-filtered using Lipinski's Rule of Five, Veber's oral "
        "bioavailability criteria, and the PAINS A/B/C substructure catalog "
        "(RDKit). Surviving compounds were docked against the target receptor "
        "using AutoDock-Vina, with GNINA CNN rescoring applied where trained "
        "weights were supplied. Top hits were profiled with ADMETlab 3.0 "
        f"({admet_source} used for this run)."
    )

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
